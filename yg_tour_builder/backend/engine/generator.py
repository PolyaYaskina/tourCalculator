from __future__ import annotations

from typing import Dict, Any, List
from io import BytesIO

from docx import Document
import pandas as pd


def create_word(days: Dict[int, Dict[str, Any]]) -> bytes:
    """Generate a Word document representing the itinerary."""
    doc = Document()
    for day in sorted(days):
        info = days[day]
        doc.add_heading(f"Day {day}", level=1)
        if info.get("description"):
            doc.add_paragraph(info["description"])
        for service in info.get("services", []):
            doc.add_paragraph(service, style="List Bullet")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()



def create_excel(detail: List[Dict[str, Any]]) -> bytes:
        buffer = BytesIO()
        workbook = pd.ExcelWriter(buffer, engine="xlsxwriter")
        wb = workbook.book
        ws = workbook.book.add_worksheet("Смета")
        workbook.sheets["Смета"] = ws

        # 🔧 Форматы
        header_format = wb.add_format({"bold": True, "bg_color": "#D9D9D9", "border": 1})
        day_format = wb.add_format({"bold": True, "bg_color": "#F4F4F4", "border": 1})
        total_format = wb.add_format({"bold": True, "bg_color": "#FFF2CC", "border": 1})
        border_format = wb.add_format({"border": 1})

        # 📋 Заголовки
        headers = ["День", "Опция", "Цена", "Кол-во", "Сумма", "Сумма с НДС"]
        for col, name in enumerate(headers):
            ws.write(0, col, name, header_format)
            ws.set_column(col, col, 18)

        row = 1
        total = 0

        # 📊 Группировка по дню
        grouped = {}
        for entry in detail:
            grouped.setdefault(entry["day"], []).append(entry)

        for day, rows in grouped.items():
            # 🗓 Заголовок дня
            ws.merge_range(row, 0, row, 5, f"🗓 День {day}", day_format)
            row += 1
            for r in rows:
                ws.write(row, 0, r["day"], border_format)
                ws.write(row, 1, r["service"], border_format)
                ws.write(row, 2, r["price"], border_format)
                ws.write(row, 3, r["qty"], border_format)
                ws.write(row, 4, r["sum"], border_format)
                ws.write(row, 5, r["sumWithNDS"], border_format)
                total += r["sumWithNDS"]
                row += 1

        # 📦 Итоговая строка
        ws.merge_range(row, 0, row, 4, "ИТОГО", total_format)
        ws.write(row, 5, total, total_format)

        workbook.close()
        buffer.seek(0)
        return buffer.getvalue()
